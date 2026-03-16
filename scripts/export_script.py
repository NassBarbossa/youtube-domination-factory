"""Export YouTube script to .docx for Google Docs upload."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def add_styled_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_marker(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x6C, 0x5C, 0xE7)
    run.font.size = Pt(10)

def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if italic:
        run.italic = True

def add_text_on_screen(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[TEXT ON SCREEN] {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xE1, 0x7A, 0x00)
    run.font.size = Pt(10)

def build_script(output_path):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('10 Tips Pour Mieux Utiliser Claude Code', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    run = p.add_run('Format: Tutorial | Durée cible: ~12 min | Langue: Français')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x63, 0x63, 0x63)
    doc.add_paragraph('─' * 50)

    # HOOK
    add_styled_heading(doc, 'HOOK (0:00 - 0:10)', level=2)
    add_marker(doc, '[FACE CAM]')
    add_body(doc, "Tu utilises Claude Code tous les jours. Mais je te garantis que t'exploites même pas 30% de ce qu'il sait faire.")
    add_body(doc, "(léger sourire)", italic=True)
    add_body(doc, "Aujourd'hui je te donne 10 tips qui vont changer ta façon de bosser avec.")

    # CONTEXT
    add_styled_heading(doc, 'CONTEXT (0:10 - 0:30)', level=2)
    add_marker(doc, '[FACE CAM]')
    add_body(doc, "J'utilise Claude Code depuis le premier jour. Et au fil du temps, j'ai découvert des trucs que personne montre, parce que la plupart des gens se contentent du basique. Chaque tip que je vais te donner, c'est du temps gagné. Et du temps, quand t'es entrepreneur, c'est de l'argent.")
    add_text_on_screen(doc, '"10 tips = des heures gagnées par semaine"')

    # TIP 1
    add_styled_heading(doc, 'TIP 1 : LE FICHIER CLAUDE.MD (0:30 - 1:45)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Premier tip et probablement le plus important. Le fichier CLAUDE.md. C'est un fichier que tu mets à la racine de ton projet, et Claude le lit à chaque conversation. C'est comme un brief permanent.")
    add_marker(doc, '[DEMO] Ouvrir un projet, créer un fichier CLAUDE.md à la racine')
    add_body(doc, "Tu peux y mettre ton style de code préféré, les conventions de ton projet, ce que Claude doit éviter... Tout ce que tu répètes à chaque conversation, tu le mets là une fois, et c'est réglé.")
    add_text_on_screen(doc, '"CLAUDE.md = ton brief permanent"')
    add_body(doc, "(direct)", italic=True)
    add_body(doc, "Ça a l'air simple, mais 90% des utilisateurs l'ont pas. Et ils perdent du temps à ré-expliquer les mêmes trucs à chaque session.")

    # TIP 2
    add_styled_heading(doc, 'TIP 2 : SOIS PRÉCIS DANS TES DEMANDES (1:45 - 2:50)', level=2)
    add_marker(doc, '[FACE CAM]')
    add_body(doc, 'Tip numéro 2, et c\'est du bon sens mais personne le fait. Arrête de dire "fais-moi un site web". C\'est comme si tu disais à un architecte "construis-moi une maison" sans dire combien de pièces, quel style, quel budget.')
    add_marker(doc, '[SCREEN] Montrer un mauvais prompt vs un bon prompt côte à côte')
    add_text_on_screen(doc, '❌ "Fais-moi un site web"')
    add_text_on_screen(doc, '✅ "Crée une landing page avec un hero, 3 sections features, un pricing à 2 colonnes, et un footer. Style moderne, couleurs sombres, police sans-serif."')
    add_body(doc, "Plus t'es précis, moins tu fais d'allers-retours. Et moins d'allers-retours, c'est moins de tokens brûlés et plus de résultats du premier coup.")

    # TIP 3
    add_styled_heading(doc, 'TIP 3 : UTILISE LES SLASH COMMANDS (2:50 - 4:00)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Y'a des commandes intégrées dans Claude Code que beaucoup connaissent pas.")
    add_marker(doc, '[DEMO] Ouvrir Claude Code et taper /help pour montrer la liste')
    add_body(doc, 'Par exemple, /compact — ça compresse ta conversation pour libérer du contexte sans perdre le fil. Quand t\'es sur un long projet et que Claude commence à "oublier" des trucs, c\'est parce que la fenêtre de contexte est pleine. Un petit /compact et c\'est reparti.')
    add_body(doc, "Y'a aussi /init qui te crée un CLAUDE.md automatiquement en analysant ton projet. Oui, le tip 1 peut se faire tout seul.")
    add_text_on_screen(doc, '"/compact, /init, /help — tes raccourcis"')
    add_body(doc, "(sourire)", italic=True)
    add_body(doc, "C'est les raccourcis clavier de Claude Code. Apprends-les, tu me remercieras.")

    # TIP 4
    add_styled_heading(doc, 'TIP 4 : LE MODE PLAN (4:00 - 5:10)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Tip numéro 4, et celui-là c'est un game changer. Avant de demander à Claude de coder quoi que ce soit, demande-lui d'abord de te faire un plan.")
    add_marker(doc, '[SCREEN] Taper : "Avant de coder, propose-moi un plan d\'implémentation étape par étape"')
    add_body(doc, "Pourquoi ? Parce que si tu le laisses foncer direct, il va prendre des décisions à ta place. Et parfois c'est pas les bonnes. En lui demandant un plan d'abord, tu valides la direction avant qu'il écrive la moindre ligne.")
    add_text_on_screen(doc, '"Plan d\'abord, code ensuite"')
    add_body(doc, "C'est comme un GPS. Tu vérifies l'itinéraire avant de démarrer, pas après avoir roulé 200 bornes.")

    # TIP 5
    add_styled_heading(doc, 'TIP 5 : LES AGENTS ET SOUS-AGENTS (5:10 - 6:30)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, 'Claude Code peut lancer des sous-agents. Concrètement, il peut se diviser en plusieurs "copies" de lui-même qui bossent en parallèle.')
    add_marker(doc, '[DEMO] Montrer un cas concret : demander une recherche qui lance un sous-agent')
    add_body(doc, "Imagine : tu lui demandes de refactorer un projet. Au lieu de tout faire à la suite, il peut envoyer un agent chercher les fichiers concernés pendant qu'un autre prépare la structure. C'est du travail parallèle.")
    add_body(doc, "Et depuis peu, y'a les Agent Teams — là c'est encore un cran au-dessus. Plusieurs instances de Claude Code qui bossent ensemble, se parlent, et se partagent les tâches.")
    add_text_on_screen(doc, '"Sous-agents = travail parallèle = plus rapide"')
    add_body(doc, "(posé)", italic=True)
    add_body(doc, "Pense à ça comme une équipe. Sauf que l'équipe coûte quelques centimes de tokens.")

    # TIP 6
    add_styled_heading(doc, 'TIP 6 : DONNE-LUI DU CONTEXTE FICHIER (6:30 - 7:30)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Quand tu bosses sur un projet, Claude Code peut lire tes fichiers. Mais le truc que beaucoup savent pas, c'est que tu peux le diriger vers les bons fichiers.")
    add_marker(doc, '[SCREEN] Montrer : "Lis le fichier src/components/Header.tsx et améliore le responsive"')
    add_body(doc, 'Au lieu de dire "améliore mon header", dis-lui quel fichier regarder. Ça évite qu\'il cherche partout, qu\'il perde du contexte, et qu\'il modifie le mauvais fichier.')
    add_text_on_screen(doc, '"Pointe vers le fichier → résultat précis"')
    add_body(doc, "C'est comme donner une adresse exacte au lieu de dire \"c'est quelque part dans Paris\". Ça aide.")

    # TIP 7
    add_styled_heading(doc, 'TIP 7 : LES SKILLS PERSONNALISÉS (7:30 - 8:50)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Celle-là c'est ma préférée. Tu peux créer des skills — des instructions packagées dans un dossier — que Claude charge automatiquement quand il en a besoin.")
    add_marker(doc, '[DEMO] Montrer la structure d\'un skill : dossier + SKILL.md')
    add_body(doc, "Par exemple, moi j'ai un skill qui écrit mes scripts YouTube. Un autre qui génère mes titres SEO. Quand je dis \"écris le script\", Claude sait exactement comment je veux que ce soit fait, mon style, mon audience, mon format.")
    add_text_on_screen(doc, '"Skill = tu formes Claude une fois, il applique à vie"')
    add_body(doc, "(enthousiaste mais posé)", italic=True)
    add_body(doc, "C'est la différence entre un assistant générique et un assistant qui te connaît. Et ça, pour un entrepreneur, ça vaut de l'or.")

    # TIP 8
    add_styled_heading(doc, 'TIP 8 : ITÈRE, CORRIGE, AFFINE (8:50 - 9:45)', level=2)
    add_marker(doc, '[FACE CAM]')
    add_body(doc, "Erreur classique : Claude te sort un résultat, c'est pas parfait, et tu recommences de zéro. Non. Itère.")
    add_body(doc, 'Dis-lui ce qui va pas. "Le ton est trop formel", "raccourcis cette partie", "ajoute un exemple concret ici". Claude apprend dans la conversation. Plus tu le corriges, plus il s\'ajuste.')
    add_text_on_screen(doc, '"Ne recommence jamais de zéro — itère"')
    add_body(doc, "(direct)", italic=True)
    add_body(doc, "Pense à Claude comme un collaborateur junior ultra rapide. Tu le briefes, il livre, tu corriges, il améliore. Trois allers-retours et t'as un résultat que t'aurais mis des heures à faire toi-même.")

    # TIP 9
    add_styled_heading(doc, 'TIP 9 : UTILISE LA MÉMOIRE (9:45 - 10:50)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Claude Code a une mémoire persistante. Dans le dossier .claude/memory/, il peut stocker des infos qui survivent entre les conversations.")
    add_marker(doc, '[SCREEN] Montrer le dossier memory et un fichier MEMORY.md')
    add_body(doc, "Tes préférences, tes patterns récurrents, les solutions à des problèmes que t'as déjà résolus. Au lieu de tout ré-expliquer à chaque session, Claude se souvient.")
    add_text_on_screen(doc, '"Mémoire persistante = un assistant qui apprend"')
    add_body(doc, "C'est comme la différence entre un nouveau stagiaire chaque lundi et un employé qui est là depuis 6 mois. L'un a besoin qu'on lui répète tout, l'autre sait déjà comment tu fonctionnes.")

    # TIP 10
    add_styled_heading(doc, 'TIP 10 : AUTOMATISE TES WORKFLOWS (10:50 - 12:00)', level=2)
    add_marker(doc, '[FACE CAM + SCREEN]')
    add_body(doc, "Dernier tip et le plus puissant. Combine tout ce qu'on vient de voir. CLAUDE.md + Skills + Agents + Mémoire = un workflow automatisé complet.")
    add_marker(doc, '[DEMO] Montrer un exemple : lancer un workflow qui enchaîne plusieurs étapes automatiquement')
    add_body(doc, "Moi par exemple, quand je prépare une vidéo YouTube, Claude Code me génère le script, les titres SEO, la description, le brief miniature — le tout en parallèle avec des agents. Ce que je faisais en 3 heures, c'est fait en 20 minutes.")
    add_text_on_screen(doc, '"Le vrai pouvoir = combiner les tips ensemble"')
    add_body(doc, "(serious)", italic=True)
    add_body(doc, "Et c'est ça le vrai game. C'est pas un outil que t'utilises de temps en temps. C'est un système que tu construis une fois et qui bosse pour toi à chaque fois.")

    # CTA
    add_styled_heading(doc, 'CTA + OUTRO (12:00 - 12:30)', level=2)
    add_marker(doc, '[FACE CAM]')
    add_body(doc, "Si t'as appris au moins un truc dans cette vidéo, dis-moi en commentaire c'est lequel. Et si tu veux que je fasse un deep dive sur un de ces 10 tips, dis-le moi aussi — je ferai une vidéo dédiée.")
    add_body(doc, "(léger sourire)", italic=True)
    add_body(doc, "Et si t'es pas encore abonné... tu sais quoi faire. À la prochaine.")

    # Metadata
    doc.add_paragraph('─' * 50)
    add_styled_heading(doc, 'Metadata', level=2)
    add_body(doc, 'Word count: ~1800 mots')
    add_body(doc, 'Estimated reading time: ~12 min')
    add_body(doc, 'Timestamps:')
    timestamps = [
        '0:00 - Intro',
        '0:30 - Tip 1 : CLAUDE.md',
        '1:45 - Tip 2 : Sois précis',
        '2:50 - Tip 3 : Slash commands',
        '4:00 - Tip 4 : Mode plan',
        '5:10 - Tip 5 : Agents & sous-agents',
        '6:30 - Tip 6 : Contexte fichier',
        '7:30 - Tip 7 : Skills personnalisés',
        '8:50 - Tip 8 : Itère et affine',
        '9:45 - Tip 9 : La mémoire',
        '10:50 - Tip 10 : Automatise tout',
    ]
    for ts in timestamps:
        doc.add_paragraph(ts, style='List Bullet')

    doc.save(output_path)
    print(f"Script exported to: {output_path}")

if __name__ == '__main__':
    output = sys.argv[1] if len(sys.argv) > 1 else os.path.expanduser('~/youtube-domination-factory/exports/10-tips-claude-code.docx')
    os.makedirs(os.path.dirname(output), exist_ok=True)
    build_script(output)
